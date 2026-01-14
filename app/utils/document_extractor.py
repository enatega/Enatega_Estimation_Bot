import os
import pdfplumber
import PyPDF2
from docx import Document
from typing import List, Dict
import logging

logger = logging.getLogger(__name__)

class DocumentExtractor:
    """Extract text content from PDF and DOCX files"""
    
    def __init__(self, data_dir: str = "."):
        self.data_dir = data_dir
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file using pdfplumber (better for tables)"""
        # Handle both absolute and relative paths
        if os.path.isabs(file_path):
            full_path = file_path
        else:
            full_path = os.path.join(self.data_dir, file_path)
        text_content = ""
        
        try:
            with pdfplumber.open(full_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                with open(full_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n\n"
            except Exception as e2:
                logger.error(f"Failed to extract from {file_path}: {e2}")
        
        return text_content.strip()
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        # Handle both absolute and relative paths
        if os.path.isabs(file_path):
            full_path = file_path
        else:
            full_path = os.path.join(self.data_dir, file_path)
        text_content = ""
        
        try:
            doc = Document(full_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content += row_text + "\n"
                text_content += "\n"
        except Exception as e:
            logger.error(f"Failed to extract from {file_path}: {e}")
        
        return text_content.strip()
    
    def extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file"""
        # Handle both absolute and relative paths
        if os.path.isabs(file_path):
            full_path = file_path
        else:
            full_path = os.path.join(self.data_dir, file_path)
        text_content = ""
        
        try:
            with open(full_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
        except Exception as e:
            logger.error(f"Failed to extract from {file_path}: {e}")
        
        return text_content.strip()
    
    def extract_all_documents(self) -> Dict[str, str]:
        """Extract text from all documents"""
        extracted_data = {}
        
        # Extract TXT files (feature estimates)
        for txt_file in ["Estimates.txt"]:
            if os.path.exists(os.path.join(self.data_dir, txt_file)):
                logger.info(f"Extracting {txt_file}...")
                extracted_data[txt_file] = self.extract_txt_text(txt_file)
        
        # Extract PDFs (excluding Estimation Calculator Data - using Estimates.txt instead)
        for pdf_file in [
            "content (3).pdf",
            "content (4).pdf",
            "content (5).pdf",
            "content (6).pdf",
            "content (7).pdf",
            "content (8).pdf"
        ]:
            if os.path.exists(os.path.join(self.data_dir, pdf_file)):
                logger.info(f"Extracting {pdf_file}...")
                extracted_data[pdf_file] = self.extract_pdf_text(pdf_file)
        
        # Extract DOCX
        for docx_file in ["Enatega_Product_Overview.docx"]:
            if os.path.exists(os.path.join(self.data_dir, docx_file)):
                logger.info(f"Extracting {docx_file}...")
                extracted_data[docx_file] = self.extract_docx_text(docx_file)
        
        return extracted_data
    
    def get_chatgpt_examples(self, extracted_data: Dict[str, str]) -> str:
        """Get Estimates.txt as primary reference (replaces Estimation Calculator Data)"""
        estimates_file = "Estimates.txt"
        if estimates_file in extracted_data:
            return extracted_data[estimates_file]
        return ""
